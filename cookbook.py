from json import loads # From Json, import the loads module to dump the info where it needs to be
from requests import get # From requests import get module
import docx  # Import the docx module so that we can create a word document
from docx.shared import Inches
def get_taco():
    '''Gets random taco from the API'''
    r = get('https://taco-1150.herokuapp.com/random/?full_taco=true')
    taco = loads(r.text)
    return taco
def get_shell(taco):
    '''Prints and returns the taco shell'''
    print(f"Shell: {taco['shell']['name']}")
    return taco['shell']
def get_base_layer(taco):
    '''Prints and returns the taco's base layer'''
    print(f"Base layer: {taco['base_layer']['name']}")
    return taco['base_layer']
def get_mixin(taco):
    '''Prints and returns what you mix in with the base layer'''
    print(f"Mixin: {taco['mixin']['name']}")
    return taco['mixin']
def get_condiment(taco):
    '''Prints and returns the condiment to use'''
    print(f"Condiment: {taco['condiment']['name']}")
    return taco['condiment']
def get_seasoning(taco):
    '''Prints and returns the seasoning to use'''
    print(f"Seasoning: {taco['seasoning']['name']}")
    return taco['seasoning']
def full_recipe(taco):
    '''Writes full recipe
       and puts it in a word document'''
    shell = taco['shell']['name']
    main = taco['base_layer']['name']
    mixin = taco['mixin']['name']
    c = taco['condiment']['name']
    s = taco['seasoning']['name']
    title = f'# {main} & {mixin} with {c} and {s} in/in a {shell}'
    recipe.add_paragraph(title)
    recipe.add_paragraph(f"\n\n## {taco['shell']['recipe']}")
    recipe.add_paragraph(f"\n\n## {taco['base_layer']['recipe']}")
    recipe.add_paragraph(f"\n\n## {taco['mixin']['recipe']}")
    recipe.add_paragraph(f"\n\n## {taco['condiment']['recipe']}")
    recipe.add_paragraph(f"\n\n## {taco['seasoning']['recipe']}")
    recipe.save('Cookbook.docx')
'''The first page heading, added pic, and author name, and url'''
recipe = docx.Document()
recipe.add_heading('Random Taco Cookbook', 0)
recipe.add_picture('Random_Taco_Cookbook.jpg', width=Inches(6.25))
recipe.add_paragraph('Miguel Andrade')
recipe.add_paragraph('https://taco-1150.herokuapp.com/random/?full_taco=true')
'''Coordinates the ingredients with the recipe for each recipe and calls it from their definitions'''
if __name__ == '__main__':
    taco = get_taco()
    shell = get_shell(taco)
    base_layer = get_base_layer(taco)
    mixin = get_mixin(taco)
    condiment = get_condiment(taco)
    seasoning = get_seasoning(taco)
    recipe = full_recipe(taco)
